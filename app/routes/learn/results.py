import os
from docx import Document
from docx.table import _Cell

from app.auth import auth
from app.api import bp
from flask import request
from flask import current_app
from app.tag_model import Tag
from app.models.user import User
from app.result_model import Result
from app.subject_model import Subject

@bp.route('/cop')
def cop():
    base_path = os.path.join(current_app.instance_path, 'prenursery')
    for user in User.query.filter_by(grade='pre-nursery'):
        path = os.path.join(base_path, user.username) + '.docx'
        try:
            document = Document(path)
            for table in document.tables:
                for row in table.rows:
                    row_cells = [_Cell(tc, table) for tc in row._tr.tc_lst]
                    for cell in row_cells:
                        subject = Subject.query.filter_by(label=cell.text).first()
                        if subject:
                            print(user.username, subject.label)
                            index = row_cells.index(cell) + 1
                            score = row_cells[index].text
                            Result(user, subject, score)
        except Exception as e:
            print(e)
    return '202'

@bp.route('/copy')
def copy():
    base_path = os.path.join(current_app.instance_path, 'nur2_second')
    for user in User.query.filter(User.grade.contains('nursery')):
        path = os.path.join(base_path, user.username) + '.docx'
        try:
            document = Document(path)
            for table in document.tables:
                for row in table.rows:
                    row_cells = [_Cell(tc, table) for tc in row._tr.tc_lst]
                    for cell in row_cells:
                        subject = Subject.query.filter_by(label=cell.text).first()
                        if subject:
                            print(user.username, subject.label)
                            index = row_cells.index(cell) + 1
                            score = row_cells[index].text
                            Result(user, subject, score)
        except Exception as e:
            print(e)
    return '202'  

@bp.route('/results', methods=['GET'])
@auth
def user_result(user=None):
    headers = [
            {'key': 'subject', 'value': 'Subject'},
            {'key': 'score', 'value': 'Score'}
        ]
    tables = []
    for tag in Tag.query:
        rows = []
        for subject in tag.subjects:
            results = Result.query.filter_by(user_id=user.id).filter_by(subject_id=subject.id)
            for r in results:
                rows.append(r.dict())
        table = {'title': tag.name, 'rows': rows}
        tables.append(table)
    return {'tables': tables, 'headers': headers}